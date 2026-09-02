# -*- coding: utf-8 -*-
"""
文档智能脱敏工具 v8
"""

import os, re, json, sys, threading
from datetime import datetime

os.environ["QT_QPA_PLATFORM"] = "windows"

from PySide6.QtCore import Qt, Signal, QObject, QMetaObject, Q_ARG
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QRadioButton, QButtonGroup, QFileDialog,
    QTableWidget, QTableWidgetItem, QProgressBar, QPlainTextEdit,
    QLineEdit, QMessageBox, QGroupBox, QHeaderView, QTextEdit,
    QSplitter, QCheckBox, QTabWidget, QComboBox, QDialog,
    QMenu, QStyleFactory, QGridLayout, QListWidget, QListWidgetItem,
    QAbstractItemView, QMenuBar,
)
from PySide6.QtGui import QAction, QShortcut, QKeySequence

NL = chr(10)
MASK = chr(0x3010) + chr(0x25A0) + chr(0x25A0) + chr(0x25A0) + chr(0x3011)

# ============================================================
# 应用常量（补丁 #2）
# ============================================================
APP_NAME     = "脱敏马"
APP_VERSION  = "V1.8.1"
APP_AUTHOR   = "脱敏马马"
UPDATE_URL   = "https://your-server.com/tuominma/update.json"

# ============================================================
# 配置文件路径（补丁 #1）
# ============================================================
def _get_config_dir():
    """返回配置目录路径，优先使用 EXE 同级目录"""
    if getattr(sys, 'frozen', False):
        base = os.path.dirname(sys.executable)
    else:
        base = os.path.dirname(os.path.abspath(__file__))
    cfg_dir = os.path.join(base, "config")
    os.makedirs(cfg_dir, exist_ok=True)
    return cfg_dir

CONFIG_FILE = os.path.join(_get_config_dir(), "tuomin_config.json")

STYLESHEET = """
QWidget {
    font-family: "Microsoft YaHei", "Segoe UI", "PingFang SC", sans-serif;
    font-size: 9pt;
    color: #2C3E50;
}
QMainWindow { background: #F5F6FA; }
QGroupBox {
    background: #FFFFFF;
    border: 1px solid #E0E4E8;
    border-radius: 6px;
    margin-top: 12px;
    padding: 16px 12px 12px 12px;
    font-weight: bold;
}
QGroupBox::title {
    subcontrol-origin: margin;
    subcontrol-position: top left;
    padding: 2px 10px;
    background: #FFFFFF;
    border: 1px solid #E0E4E8;
    border-bottom: none;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    color: #4A90D9;
    font-weight: bold;
}
QPushButton {
    background: #4A90D9;
    color: white;
    border: none;
    border-radius: 4px;
    padding: 6px 16px;
    min-height: 28px;
    font-weight: bold;
}
QPushButton:hover { background: #357ABD; }
QPushButton:pressed { background: #2C6BA0; }
QPushButton:disabled { background: #BDC3C7; }
QPushButton#btn_process {
    background: #27AE60;
    font-size: 11pt;
    padding: 8px 24px;
    min-height: 36px;
}
QPushButton#btn_process:hover { background: #219A52; }
QPushButton#btn_add, QPushButton#btn_clear,
QPushButton#btn_choose_dir, QPushButton#btn_open_output {
    background: #ECF0F1;
    color: #2C3E50;
    border: 1px solid #D5D8DC;
}
QPushButton#btn_add:hover, QPushButton#btn_clear:hover,
QPushButton#btn_choose_dir:hover, QPushButton#btn_open_output:hover {
    background: #D5D8DC;
}
QPushButton#btn_wl_add, QPushButton#btn_bl_add {
    background: #27AE60;
    min-width: 32px;
    padding: 4px 12px;
}
QPushButton#btn_wl_add:hover, QPushButton#btn_bl_add:hover { background: #219A52; }
QPushButton#btn_wl_del, QPushButton#btn_bl_del {
    background: #E74C3C;
    min-width: 32px;
    padding: 4px 12px;
}
QPushButton#btn_wl_del:hover, QPushButton#btn_bl_del:hover { background: #C0392B; }
QLineEdit {
    border: 1px solid #D5D8DC;
    border-radius: 4px;
    padding: 4px 8px;
    min-height: 24px;
    background: #FFFFFF;
}
QLineEdit:focus { border-color: #4A90D9; }
QTableWidget {
    background: #FFFFFF;
    border: 1px solid #E0E4E8;
    border-radius: 4px;
    gridline-color: #ECF0F1;
    selection-background-color: #D6EAF8;
    selection-color: #2C3E50;
}
QTableWidget::item { padding: 4px 8px; }
QHeaderView::section {
    background: #2C3E50;
    color: white;
    padding: 6px 8px;
    border: none;
    font-weight: bold;
}
QPlainTextEdit, QTextEdit {
    background: #FFFFFF;
    border: 1px solid #E0E4E8;
    border-radius: 4px;
    padding: 4px;
}
QPlainTextEdit:focus, QTextEdit:focus { border-color: #4A90D9; }
QProgressBar {
    background: #ECF0F1;
    border: none;
    border-radius: 4px;
    text-align: center;
    min-height: 20px;
    color: #2C3E50;
    font-weight: bold;
}
QProgressBar::chunk {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #4A90D9, stop:1 #27AE60);
    border-radius: 4px;
}
QTabWidget::pane {
    background: #FFFFFF;
    border: 1px solid #E0E4E8;
    border-radius: 4px;
    padding: 8px;
}
QTabBar::tab {
    background: #ECF0F1;
    color: #7F8C8D;
    border: 1px solid #E0E4E8;
    border-bottom: none;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    padding: 8px 16px;
    margin-right: 2px;
    font-weight: bold;
}
QTabBar::tab:selected {
    background: #FFFFFF;
    color: #4A90D9;
    border-bottom: 2px solid #4A90D9;
}
QTabBar::tab:hover { color: #2C3E50; }
QRadioButton { spacing: 6px; }
QRadioButton::indicator {
    width: 16px; height: 16px;
    border-radius: 8px;
    border: 2px solid #BDC3C7;
}
QRadioButton::indicator:checked {
    background: #4A90D9;
    border-color: #4A90D9;
}
QCheckBox { spacing: 6px; padding: 2px 0; }
QCheckBox::indicator {
    width: 16px; height: 16px;
    border-radius: 3px;
    border: 2px solid #BDC3C7;
}
QCheckBox::indicator:checked {
    background: #4A90D9;
    border-color: #4A90D9;
}
QComboBox {
    border: 1px solid #D5D8DC;
    border-radius: 4px;
    padding: 4px 8px;
    min-height: 24px;
    background: #FFFFFF;
}
QComboBox:focus { border-color: #4A90D9; }
QListWidget {
    background: #FFFFFF;
    border: 1px solid #E0E4E8;
    border-radius: 4px;
    padding: 2px;
}
QListWidget::item {
    padding: 4px 8px;
    border-radius: 3px;
}
QListWidget::item:selected {
    background: #D6EAF8;
    color: #2C3E50;
}
QScrollBar:vertical {
    background: #F5F6FA;
    width: 8px;
    border-radius: 4px;
}
QScrollBar::handle:vertical {
    background: #BDC3C7;
    border-radius: 4px;
    min-height: 30px;
}
QScrollBar::handle:vertical:hover { background: #95A5A6; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
"""


class AppConfig:
    def __init__(self):
        self.custom_whitelist = []
        self.custom_blacklist = []
        self.enabled_types = {
            "人名": True, "公司名": True, "电话": True, "邮箱": True,
            "地址": True, "身份证": True, "银行账号": True, "信用代码": True,
            "合同编号": True, "邮编": True, "SWIFT": True, "网址": True, "自定义词条": True,
        }
        self.sensitivity = "标准"
        self.last_output_dir = ""
        self.load()

    def load(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                self.custom_whitelist = data.get("custom_whitelist", [])
                self.custom_blacklist = data.get("custom_blacklist", [])
                self.enabled_types.update(data.get("enabled_types", {}))
                self.sensitivity = data.get("sensitivity", "标准")
                self.last_output_dir = data.get("last_output_dir", "")
            except Exception:
                pass

    def save(self):
        try:
            data = {
                "custom_whitelist": self.custom_whitelist,
                "custom_blacklist": self.custom_blacklist,
                "enabled_types": self.enabled_types,
                "sensitivity": self.sensitivity,
                "last_output_dir": self.last_output_dir,
            }
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass


class MappingManager:
    def __init__(self):
        self._map = {}
        self._rev = {}

    def add(self, original, placeholder):
        self._map[original] = placeholder
        self._rev[placeholder] = original

    def clear(self):
        self._map.clear()
        self._rev.clear()

    @property
    def count(self):
        return len(self._map)

    def get_all_reverse(self):
        return dict(self._rev)

    def get_all_forward(self):
        return dict(self._map)

    def save_to_file(self, path):
        with open(path, "w", encoding="utf-8") as f:
            json.dump({"forward": self._map, "reverse": self._rev}, f, ensure_ascii=False, indent=2)

    @classmethod
    def load_from_file(cls, path):
        mm = cls()
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            mm._map = data.get("forward", {})
            mm._rev = data.get("reverse", {})
        return mm


COMMON_TERMS = set(
    "响应人 响应方 响应 响应文件 响应报价 响应性 响应时间 响应周期 "
    "响应人名称 响应人地址 响应人代表 响应人资格 响应人资质 响应人条件 "
    "采购 采购人 采购代理 采购文件 采购内容 采购需求 采购方式 采购预算 "
    "采购项目 采购公告 采购邀请 采购合同 采购程序 采购活动 "
    "招标 招标人 招标文件 招标公告 招标代理 招标投标 招标采购 "
    "投标 投标人 投标文件 投标报价 投标保证金 投标有效期 投标截止 "
    "开标 评标 定标 中标 中标人 中标通知书 中标结果 中标金额 "
    "竞争 竞争性谈判 竞争谈判 竞争性磋商 竞争磋商 "
    "谈判 谈判小组 谈判文件 谈判记录 谈判过程 谈判结果 "
    "磋商 磋商小组 磋商文件 磋商记录 磋商过程 磋商结果 "
    "询价 询价文件 询价公告 询价结果 "
    "评审 评审专家 评审委员会 评审小组 评审标准 评审办法 评审因素 "
    "评审程序 评审结果 评审报告 评审意见 评审结论 "
    "成交 成交人 成交供应商 成交价 成交金额 成交结果 成交通知书 "
    "成交公告 成交确认 成交条件 成交价格 "
    "报价 报价表 报价文件 报价清单 报价有效期 报价截止 "
    "供应商 供应商名称 供应商地址 供应商资质 供应商资格 "
    "候选 候选人 候选供应商 候选成交人 候选名单 "
    "资格 资质 资格审查 资格条件 资格要求 资格证明 资格文件 "
    "合同 合同条款 合同条件 合同内容 合同金额 合同价格 合同期限 "
    "合同实施 合同管理 合同签订 合同签署 合同生效 合同变更 合同终止 合同解除 "
    "协议 协议书 补充协议 框架协议 战略协议 合作协议 "
    "条款 条件 要求 规定 约定 规则 办法 细则 方案 "
    "甲方 乙方 丙方 丁方 买方 卖方 出租方 承租方 出租人 承租人 "
    "发包人 承包人 委托人 受托人 代理人 被代理人 "
    "债务人 债权人 保证人 担保人 抵押人 质权人 出质人 "
    "当事人 第三人 第三方 相关方 利益相关方 各方 双方 单方 "
    "总则 分则 附则 附件 附表 附录 目录 索引 "
    "第一条 第二条 第三条 第四条 第五条 第六条 第七条 第八条 第九条 第十条 "
    "第十一条 第十二条 第十三条 第十四条 第十五条 第十六条 第十七条 第十八条 第十九条 第二十条 "
    "第一款 第二款 第三款 第四款 第五款 第六款 第一项 第二项 第三项 第四项 "
    "不可抗力 意外事件 情势变更 商业风险 政策变化 "
    "违约 违约金 违约方 违约行为 违约责任 "
    "赔偿 赔偿金 赔偿损失 损害赔偿 损失赔偿 "
    "补偿 补偿金 经济补偿 损失补偿 "
    "救济 免责 豁免 免除 减轻 "
    "保密 保密义务 保密期限 保密范围 保密信息 "
    "知识产权 专有技术 技术秘密 商业秘密 技术资料 技术文档 "
    "保证 担保 抵押 质押 留置 定金 保证金 押金 质保金 "
    "保修 质保 质保期 保修期 质保金 保修金 "
    "验收 检验 检测 试验 调试 试运行 试车 "
    "交付 交货 发货 收货 签收 验收 确认 "
    "运输 搬运 装卸 包装 仓储 保管 存储 "
    "安装 组装 装配 调试 维护 保养 维修 修理 "
    "培训 指导 技术支持 技术服务 售后服务 "
    "支付 付款 收款 结算 清算 转账 汇款 托收 承付 "
    "发票 收据 凭证 单据 账单 清单 明细 清单 "
    "审计 评估 鉴定 公证 见证 认证 认可 备案 登记 "
    "生效 失效 终止 解除 变更 修改 补充 续签 展期 "
    "通知 送达 告知 确认 同意 认可 批准 核准 审批 "
    "争议 纠纷 诉讼 仲裁 调解 和解 协商 谈判 管辖 "
    "法律适用 适用法律 管辖法律 准据法 "
    "送达地址 联系地址 通讯地址 办公地址 注册地址 "
    "法定代表人 委托代理人 授权代表 授权委托人 联系人 经办人 负责人 "
    "权利 义务 责任 权力 权限 范围 期限 时效 "
    "工作日 日历日 自然日 营业日 节假 日 "
    "金额 总价 单价 合价 合计 总计 共计 小计 累计 总额 "
    "大写 小写 数字 中文 阿拉伯 罗马 大写金额 小写金额 "
    "人民币 美元 欧元 英镑 日元 港币 澳元 加元 瑞士法郎 "
    "元 角 分 整 万元 亿元 万元整 亿元整 "
    "含税 不含税 税前 税后 税率 税额 税种 税收 税务 "
    "增值税 所得税 企业所得税 个人所得税 营业税 消费税 关税 "
    "百分之 千分之 万分之 百分比 百分点 比率 比例 系数 "
    "年 月 日 周 季 季度 半年 全年 本期 上期 下期 当期 期末 期初 "
    "年利率 月利率 日利率 利率 利息 本息 本金 "
    "公斤 吨 千克 克 毫克 升 毫升 立方米 平方米 米 厘米 毫米 公里 "
    "个 只 条 件 套 批 组 箱 包 袋 桶 瓶 罐 盒 台 辆 架 艘 "
    "国家标准 行业标准 企业标准 地方标准 团体标准 国际标准 国外标准 "
    "技术规范 技术条件 技术参数 技术要求 技术指标 技术标准 技术规格 "
    "规格 型号 品牌 产地 制造商 生产商 供应商 经销商 代理商 "
    "正本 副本 原件 复印件 扫描件 电子版 纸质版 传真件 "
    "签字 盖章 签章 签署 签收 签认 签批 签发 "
    "注册资本 实收资本 总资产 净资产 营业收入 净利润 毛利润 "
    "目录 总则 范围 定义 解释 适用 "
    "公告 邀请 邀请函 通知 致 各 各供应商 各响应人 "
    "编制 递交 提交 报送 发送 接收 收取 "
    "加密 解密 密码 密钥 数字证书 电子签名 "
    "确定 确认 选定 选择 推荐 排序 排名 "
    "签约 签订 合同签订 协议签订 合同签署 协议签署 "
    "采购内容 采购范围 采购清单 采购数量 采购规格 "
    "合同资料表 合同条款 合同条件 合同格式 合同模板 "
    "报价表 报价单 报价函 报价信 报价书 "
    "响应函 响应书 响应文件 响应方案 响应计划 "
    "资格审查 资格预审 资格后审 资格条件 资格要求 "
    "符合性 响应性 实质性 形式性 资格性 技术性 商务性 "
    "偏差 偏离 修改 调整 变更 澄清 说明 补正 "
    "撤回 撤销 取消 放弃 拒绝 接受 同意 "
    "澄清 说明 补正 修正 更改 替换 "
    "现场 实地 书面 口头 电子 线上 线下 网上 网下 "
    "密封 开启 公开 公示 公告 公布 发布 "
    "监督管理 监督检查 监督 管理 审查 审核 复核 "
    "投诉 质疑 异议 申诉 举报 反映 "
    "暂停 中止 终止 停止 继续 恢复 重启 "
    "有效 无效 生效 失效 过期 逾期 到期 届满 "
    "工作日 自然日 日历日 营业日 节假日 休息日 "
    "上午 下午 中午 晚间 凌晨 早晨 傍晚 午夜 "
    "前 后 内 外 上 下 左 右 中 间 旁 侧 边 "
    "编制 制定 起草 拟定 撰写 出具 提供 提交 "
    "审核 审查 审批 批准 核准备案 备案 登记 注册 "
    "分析 评估 评价 判断 决定 意见 建议 方案 "
    "报告 函 函件 文件 文书 资料 材料 信息 "
    "数据 内容 事项 事宜 项目 任务 工作 活动 "
    "情况 状态 现状 形势 趋势 动向 方向 "
    "质量 数量 规格 标准 水平 等级 档次 类别 "
    "安全 环保 健康 节能 减排 绿色 可持续 "
    "风险 成本 效益 效率 效果 成果 收益 "
    "费用 成本 开支 支出 预算 决算 核算 "
    "原始 真实 准确 完整 及时 有效 合法 合规 "
    "一般 通常 正常 特殊 异常 紧急 临时 长期 短期 "
    "董事 监事 经理 总裁 总监 主管 主任 秘书 "
    "营业执照 许可证 资质证书 资格证明 授权书 委托书 "
    "一致 一式 两份 三份 多份 原件 副本 复印件 "
    "中文 英文 中英文 双语 对照 翻译 译本 "
    "具有 享有 拥有 持有 获得 取得 得到 "
    "按照 依照 根据 依据 参照 参考 比照 对照 "
    "包括 包含 涵盖 涉及 限于 除非 除 外 "
    "以及 和 与 或 及 并 且 而 但 然而 "
    "如果 若 如 则 但 然而 且 或 以及 并 "
    "上述 以下 如下 如上 如前 其后 此前 此后 届时 当即 "
    "本文件 本协议 本合同 本附件 本条款 本约定 本条件 "
    "任何 所有 全部 一切 部分 整体 局部 个别 "
    "主要 重要 关键 核心 基本 基础 根本 本质 "
    "其他 其余 剩余 额外 附加 补充 附属 辅助 "
    "相关 关联 对应 相应 匹配 适配 适合 适用 "
    "具体 明确 清晰 详细 详尽 完整 全面 系统 "
    "必要 充分 足够 适当 合理 合规 合法 有效 "
    "最大 最小 最多 最少 最高 最低 最早 最晚 "
    "首次 再次 多次 每次 单次 累计 连续 持续 "
    "重新 再次 重复 反复 多次 另行 单独 分别 "
    "共同 一致 统一 联合 协同 合作 配合 协调 "
    "自动 手动 人工 智能 半自动 全自动 机械 电子 "
    "书面 口头 电子 线上 线下 网络 现场 远程 "
    "保证金 质保金 押金 定金 履约保证金 投标保证金 "
    "偏离表 差异表 逐项 逐条 对照 比较 "
    "补充材料 补充文件 补充说明 补充资料 "
    "资格证明 资格文件 资质证书 营业执照 许可证 "
    "书 函 函件 信 信函 表 单 证 文件 资料 "
    "目录 索引 标题 章节 页 页码 编号 "
    "见 详见 参见 参照 按照 依照 根据 "
    "报价一览表 分项报价表 报价明细表 报价汇总表 "
    "技术偏离表 商务偏离表 合同资料表 "
    "主要合同条款 采购内容 采购范围 采购清单 采购数量 采购规格 "
    "资格证明文件 资格文件 资质文件 证明文件 "
    "其他补充材料 其他材料 补充材料 补充文件 "
    "响应文件 响应函 响应书 响应方案 响应计划 "
    "竞争谈判 竞争性谈判 竞争性磋商 竞争谈判采购 "
    "确定成交 成交人 成交供应商 成交结果 "
    "响应文件编制 响应文件加密 响应文件递交 "
    "及构成 构成 组成 包含 包括 "
    "第 章 节 条 款 项 项目 附件 附表 附录 "
    .split()
)

COMMON_TERMS_EN = set(
    "Party A Party B Seller Buyer Lessor Lessee "
    "Debtor Creditor Guarantor Employer Employee Client Vendor "
    "Supplier Contractor Principal Agent "
    "This Agreement This Contract This Document This Appendix "
    "Section Article Clause Paragraph Subsection Subparagraph "
    "Effective Date Termination Date Expiration Date Due Date "
    "Signing Date Delivery Date Payment Date "
    "Force Majeure Act of God Confidential Information "
    "Intellectual Property Trade Secret Know-How "
    "Representation Warranty Indemnity Liability "
    "Arbitration Litigation Mediation Negotiation "
    "Governing Law Dispute Resolution Jurisdiction "
    "Termination Default Breach Remedy Cure "
    "Amendment Modification Supplement Addendum "
    "Exhibit Schedule Appendix Annex Attachment "
    "Notwithstanding Subject to Without prejudice "
    "Shall May Will Must Should Can Could Would "
    "Hereby Herein Hereof Hereto Hereunder Hereinafter "
    "Thereby Therein Thereof Thereto Thereunder Thereinafter "
    "Whereby Wherein Whereof Whereto "
    "Total Amount Aggregate Amount Purchase Price "
    "Consideration Deposit Penalty Liquidated Damages "
    "Payment Compensation Indemnification "
    "Tax Duty Tariff Levy Assessment "
    "Standard Specification Requirement Condition "
    "Warranty Guarantee Maintenance Service "
    "Delivery Shipment Transportation Installation "
    "Acceptance Inspection Testing Commissioning "
    "Invoice Receipt Voucher Certificate Document "
    "Original Copy Duplicate Counterpart "
    "Execute Sign Deliver Perform Fulfill "
    "Agree Accept Acknowledge Confirm Certify "
    "Notify Inform Advise Report Disclose "
    "Waive Release Relinquish Surrender Abandon "
    "Renew Extend Continue Suspend Defer "
    "Terminate Expire Cancel Revoke Rescind "
    "Amend Modify Change Revise Update "
    "Indemnify Hold Harmless Defend Protect "
    "Reimburse Compensate Pay Refund Return "
    "Shareholder Director Officer Manager Supervisor "
    "Board of Directors General Meeting "
    "Registered Capital Authorized Capital Paid-in Capital "
    "Shares Stock Equity Interest Stake "
    "Voting Right Dividend Distribution Profit Loss "
    "Asset Liability Equity Revenue Expense "
    "Audit Review Inspection Examination Verification "
    "Compliance Regulatory Supervisory Oversight "
    "Disclosure Filing Registration Approval Authorization "
    "License Permit Certificate Registration Approval "
    "Patent Trademark Copyright Design Utility Model "
    "Application Priority Date Filing Date Grant Date "
    "Licensee Licensor Assignee Assignor "
    "Royalty Fee Payment Compensation Consideration "
    "Territory Region Jurisdiction Area Field "
    "Exclusive Non-exclusive Sole Unique "
    "Sublicense Assignment Transfer Novation "
    "Term Period Duration Validity Life "
    "Renewal Extension Continuation Prolongation "
    "Termination Expiration Cancellation Revocation Rescission "
    "Breach Default Violation Infringement Contravention "
    "Remedy Cure Correction Rectification Redress "
    "Damages Compensation Indemnity Relief "
    "Injunction Specific Performance Rescission "
    "Arbitration Award Judgment Decision Order "
    "Claim Counterclaim Cross-claim "
    "Evidence Document Exhibit Schedule Attachment "
    "Witness Expert Consultant Advisor "
    "Hearing Trial Proceeding Session Conference "
    "Settlement Compromise Agreement Accord "
    "Confidentiality Non-disclosure Secrecy Privacy "
    "Proprietary Exclusive Privileged Classified "
    "Trade Secret Know-how Technical Data "
    "Software Hardware Firmware Middleware "
    "Source Code Object Code Executable Library "
    "Database Server Network System Platform "
    "Interface Protocol Standard Format Framework "
    "Upgrade Update Patch Release Version "
    "Maintenance Support Service Assistance "
    "Training Education Instruction Guidance "
    "Warranty Guarantee Undertaking Assurance "
    "Defect Error Bug Fault Issue Problem "
    "Critical Major Minor Normal Enhancement "
    "Severity Priority Status Category Type "
    "Resolution Fix Patch Workaround Solution "
    "Test Testing Inspection Check Verification "
    "Validation Certification Qualification Accreditation "
    "Quality Assurance Control Management System "
    "Risk Assessment Evaluation Analysis Review "
    "Safety Security Protection Safeguard Measure "
    "Emergency Contingency Disaster Recovery Backup "
    "Business Continuity Disaster Recovery Plan "
    "Insurance Coverage Policy Premium Deductible "
    "Limit Limit of Liability Cap Ceiling Maximum "
    "Excess Deductible Self-insured Retention "
    "Claim Notice Report File Submit "
    "Loss Damage Injury Harm Destruction "
    "Theft Fire Flood Earthquake Storm Hurricane "
    "Legal Lawful Legitimate Valid Binding Enforceable "
    "Void Voidable Null Invalid Unenforceable Illegal "
    "Fair Reasonable Equitable Just Proper "
    "Good Faith Bona Fide Fair Dealing "
    "Ordinary Reasonable Prudent Diligent Careful "
    "Material Substantial Significant Important Relevant "
    "Direct Indirect Consequential Incidental Special "
    "Punitive Exemplary Aggravated Enhanced Multiple "
    "Actual Compensatory Nominal Statutory Liquidated "
    "Specific General Special Ordinary Customary "
    "Usual Normal Standard Typical Regular "
    "Extraordinary Exceptional Unusual Unique Special "
    "Contingent Conditional Subject Dependent Based "
    "Absolute Unconditional Final Irrevocable Binding "
    "Several Joint Joint and Several "
    "Primary Secondary Tertiary Ultimate "
    "Superior Inferior Senior Junior Subordinate "
    "Principal Agent Representative Delegate Proxy "
    "Fiduciary Trustee Custodian Guardian Administrator "
    "Beneficiary Donee Grantee Successor Assignee "
    "Grantor Donor Settlor Trustor Testator "
    "Heir Legatee Devisee Next of Kin "
    "Descendant Ascendant Collateral Relative Family "
    "Spouse Husband Wife Partner Domestic Partner "
    "Child Children Minor Infant Offspring Issue "
    "Parent Father Mother Guardian Custodian "
    "Sibling Brother Sister Half-brother Half-sister "
    "Uncle Aunt Nephew Niece Cousin "
    "Grandparent Grandfather Grandmother Grandchild "
    "Ancestor Progenitor Predecessor Forefather "
    "Successor Descendant Heir Assignee Transferee "
    "Estate Trust Fund Foundation Endowment "
    "Charity Philanthropy Nonprofit Organization "
    "Corporation Company Partnership Sole Proprietorship "
    "Limited Liability Company LLC Corporation Inc "
    "Joint Venture Consortium Syndicate Pool "
    "Association Society Institute Foundation Council "
    "Board Committee Commission Task Force Working Group "
    "Member Participant Partner Shareholder Stakeholder "
    "Officer Director Manager Executive Employee "
    "Staff Personnel Worker Labor Employee Contractor "
    "Consultant Advisor Expert Specialist Professional "
    "Agent Representative Broker Dealer Trader "
    "Intermediary Middleman Mediator Arbitrator "
    "Lessor Lessee Landlord Tenant Renter "
    "Vendor Supplier Seller Provider Contractor "
    "Customer Client Buyer Purchaser Consumer "
    "Distributor Dealer Retailer Wholesaler Reseller "
    "Manufacturer Producer Maker Builder Fabricator "
    "Contractor Subcontractor Supplier Vendor "
    "Designer Engineer Architect Planner Consultant "
    "Surveyor Appraiser Assessor Evaluator Inspector "
    "Auditor Accountant Bookkeeper Controller Treasurer "
    "Lawyer Attorney Counsel Solicitor Barrister "
    "Notary Public Commissioner Justice of the Peace "
    "Judge Magistrate Justice Arbiter Referee "
    "Arbitrator Mediator Conciliator Negotiator "
    "Expert Specialist Consultant Advisor Witness "
    .split()
)

WHITELIST_WORDS = set().union(COMMON_TERMS, COMMON_TERMS_EN)

WHITELIST_PATTERNS = [
    re.compile(r"\u300a[^\u300b]{2,50}\u300b"),
    re.compile(r"\u4e2d\u534e[^\u3002]{2,40}\u6cd5"),
    re.compile(r"GB\s*[/\\-]?\s*T?\s*\d+(?:\.\d+)?(?:[-–—]\d{4})?"),
    re.compile(r"GB/T\s*\d+(?:\.\d+)?(?:[-–—]\d{4})?"),
    re.compile(r"ISO\s*\d+(?:[-–—]\d{4})?"),
    re.compile(r"QC/T\s*\d+"), re.compile(r"JB/T\s*\d+"),
    re.compile(r"HG/T\s*\d+"), re.compile(r"SH/T\s*\d+"),
    re.compile(r"SY/T\s*\d+"), re.compile(r"DL/T\s*\d+"),
    re.compile(r"NB/T\s*\d+"),
    re.compile(r"[A-Z]{2,5}[-–—]\d{3,6}[A-Z]?\d*"),
    re.compile(r"[A-Z]{1,3}\d{3,6}[A-Z]{1,3}"),
    re.compile(r"[￥¥$€£]?\s*\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\s*(?:元|万|亿|万元|亿元|美元|欧元|英镑|人民币)?"),
    re.compile(r"(?:人民币|美元|欧元|英镑|日元|港币)\s*\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?"),
    re.compile(r"\d+(?:\.\d+)?\s*[%％]"),
    re.compile(r"\d{4}年\d{1,2}月\d{1,2}日"),
    re.compile(r"\d{4}[-/.]\d{1,2}[-/.]\d{1,2}"),
    re.compile(r"\d+(?:\.\d+)?\s*[%％]\s*(?:年利率|月利率|日利率|利率)?"),
    re.compile(r"(?:年利率|月利率|日利率|利率)\s*\d+(?:\.\d+)?\s*[%％]?"),
    re.compile(r"\d+(?:\.\d+)?\s*[%％]\s*(?:税率|增值税率|所得税率)?"),
    re.compile(r"(?:税率|增值税率|所得税率)\s*\d+(?:\.\d+)?\s*[%％]?"),
    re.compile(r"\d+\s*(?:公斤|吨|千克|克|毫克|升|毫升|立方米|平方米|米|厘米|毫米|公里|个|只|条|件|套|批|组|箱|包|袋|桶|瓶|罐|盒|台|辆|架|艘|KW|kWh|kW|MW|GW|V|A|Hz|kPa|MPa|Pa|bar|psi|N|kK|Nm|rpm|L|mL|m3|m2|m|cm|mm|km|kg|g|t|ton|oz|lb|ft|in|yd|mph|kmh|°C|℃|°F)"),
]

BLACKLIST_LABELS = {
    "法定代表人": "人名", "法定代理人": "人名", "委托代理人": "人名",
    "甲方代表": "人名", "乙方代表": "人名", "授权代表": "人名",
    "授权委托人": "人名", "被授权人": "人名", "授权人": "人名",
    "经办人": "人名", "联系人": "人名", "负责人": "人名",
    "代理人": "人名", "代表人": "人名", "见证人": "人名",
    "签署人": "人名", "签字人": "人名",
    "联系电话": "电话", "电话": "电话", "手机": "电话", "手机号": "电话",
    "手机号码": "电话", "传真": "电话", "座机": "电话",
    "电子邮箱": "邮箱", "邮箱": "邮箱", "电子邮件": "邮箱",
    "邮政编码": "邮编", "邮编": "邮编",
    "地址": "地址", "通讯地址": "地址", "联系地址": "地址",
    "住所": "地址", "住所地": "地址", "注册地址": "地址", "办公地址": "地址",
    "经营地址": "地址", "实际地址": "地址",
    "网址": "网址", "网站": "网址", "官网": "网址", "主页": "网址",
    "统一社会信用代码": "信用代码", "纳税人识别号": "信用代码",
    "营业执照号": "证照号", "组织机构代码": "证照号",
    "身份证号": "身份证", "证件号码": "证照号", "证照编号": "证照号",
    "注册号": "证照号", "护照号": "证照号",
    "银行账号": "银行账号", "对公账号": "银行账号", "收款账号": "银行账号",
    "付款账号": "银行账号", "结算账号": "银行账号", "账号": "银行账号",
    "开户银行": "开户行", "开户行": "开户行",
    "银行卡号": "银行卡", "信用卡号": "银行卡",
    "SWIFT": "SWIFT", "SWIFT码": "SWIFT", "SWIFT Code": "SWIFT",
    "IBAN": "IBAN", "IBAN码": "IBAN",
    "合同编号": "合同编号", "合同号": "合同编号",
    "执行案号": "案号", "裁判文书号": "案号", "案号": "案号",
    "项目编号": "项目", "项目名称": "项目", "项目号": "项目",
    "订单编号": "订单号", "订单号": "订单号",
    "发票编号": "发票号", "发票号": "发票号",
    "Contract No": "合同编号", "Contract Number": "合同编号",
    "Agreement No": "合同编号", "Agreement Number": "合同编号",
    "Project No": "项目", "Project Name": "项目",
    "Invoice No": "发票号", "Invoice Number": "发票号",
    "Order No": "订单号", "Order Number": "订单号",
    "Case No": "案号", "Case Number": "案号", "Docket No": "案号",
    "Website": "网址", "Web": "网址", "URL": "网址", "Link": "网址",
    "Legal Representative": "人名", "Authorized Representative": "人名",
    "Contact Person": "人名", "Contact": "人名", "Witness": "人名", "Signatory": "人名",
    "Address": "地址", "Registered Address": "地址", "Office Address": "地址",
    "Phone": "电话", "Telephone": "电话", "Tel": "电话", "Mobile": "电话", "Fax": "电话",
    "Email": "邮箱", "E-mail": "邮箱",
    "Zip Code": "邮编", "Postal Code": "邮编", "Postcode": "邮编",
    "ID Number": "身份证", "ID No": "身份证", "Passport No": "证照号", "Passport Number": "证照号",
    "Tax ID": "信用代码", "Tax Identification Number": "信用代码",
    "Bank Account": "银行账号", "Account No": "银行账号", "Account Number": "银行账号",
    "Bank Account Number": "银行账号",
    "Bank Name": "开户行", "Beneficiary Bank": "开户行",
    "Credit Card No": "银行卡", "Credit Card Number": "银行卡",
    "Card No": "银行卡", "Card Number": "银行卡",
    "SWIFT Code": "SWIFT", "SWIFT BIC": "SWIFT", "BIC": "SWIFT",
}

BLACKLIST_LABELS_SORTED = sorted(BLACKLIST_LABELS.keys(), key=len, reverse=True)

COMPANY_SUFFIX = "有限公司|有限责任公司|股份有限公司|集团有限公司|集团公司|股份公司|集团|公司|银行|事务所|厂|商行|分行|支行|连锁|合作社|联合社|Inc|Corp|Corporation|LLC|Ltd|PLC|GmbH|AG|SA|SAS|SARL|BV|NV|KK|Co|Company|Limited|Incorporated|Group|Holdings|International|Technologies|Solutions|Systems|Industries|Enterprises|Ventures|Partners|Associates|Consulting|Management|Services|Capital|Investments|Laboratories|Laboratory|Institute|University|College|School|Academy|Foundation"

COMPANY_EXCLUDE_SET = set(
    "国务院 全国人大 全国政协 最高人民法院 最高人民检察院 最高法 最高检 "
    "公安部 教育部 科技部 工信部 民政部 司法部 财政部 人社部 自然资源部 "
    "生态环境部 住建部 交通运输部 水利部 农业农村部 商务部 文旅部 卫健委 "
    "人民银行 国家发改委 国家税务总局 海关总署 市场监管总局 "
    "新华社 人民日报 中央电视台 北京大学 清华大学 中国人民大学 "
    "复旦大学 上海交通大学 浙江大学 中国科学院 中国社会科学院 "
    "中国共产党 国务院办公厅 中央军委 全国总工会 全国妇联 共青团中央 "
    .split()
)

COMMON_SURNAMES = set(
    "赵钱孙李周吴郑王冯陈褚卫蒋沈韩杨朱秦尤许何吕施张孔曹严华金魏陶姜"
    "戚谢邹喻柏水窦章云苏潘葛奚范彭郎鲁韦昌马苗凤花方俞任袁柳丰鲍史唐"
    "费廉岑薛雷贺倪汤滕殷罗毕郝邬安常乐于时傅皮卞齐康伍余元卜顾孟黄"
    "和穆萧尹姚邵湛汪祁毛禹狄米贝明臧计伏成戴谈宋茅庞熊纪舒屈项祝董梁"
    "杜阮蓝闵席季麻强贾路娄危江童颜郭梅盛林刁钟徐邱骆高夏蔡田樊胡凌霍"
    "虞万支柯昝管卢莫经房裘缪干解应宗丁宣贲邓郁单杭洪包诸左石崔吉钮龚"
    "程嵇邢滑裴陆荣翁羊於惠甄曲家封芮羿储靳汲邴糜松井段富巫乌焦巴弓"
    "牧隗山谷车侯宓蓬全郗班仰秋仲伊宫宁仇栾暴甘钭厉戎祖武符刘景詹束龙"
    "叶幸司韶郜黎蓟薄印宿白怀蒲邰从鄂索咸籍赖卓蔺屠蒙池乔阴胥能苍双"
    "闻莘党翟谭贡劳逄姬申扶堵冉宰郦雍郤璩桑桂濮牛寿通边扈燕冀郏浦尚农"
    "温别庄晏柴瞿阎充慕连茹习宦艾鱼容向古易慎戈廖庾终暨居衡步耿满弘"
    "匡国文寇广禄阙东欧殳沃利蔚越夔隆师巩厍聂晁勾敖融冷訾辛阚那简饶空"
    "曾毋沙乜养鞠须巢关蒯相查后荆红游竺权逯盖益桓公万俟司马上官欧阳"
    "夏侯诸葛闻人东方赫连皇甫尉迟公羊澹台公冶宗政濮阳淳于单于太叔申屠"
    "公孙仲孙轩辕令狐钟离宇文长孙慕容鲜于闾丘司徒司空亓官司寇仉督子车"
    "颛孙端木巫马公西漆雕乐正壤驷公良拓跋夹谷宰父谷梁晋楚闫法汝鄢涂钦"
    "段干百里东郭南门呼延归海羊舌微生岳帅缑亢况后有琴梁丘左丘东门西门"
    "商牟佘佴伯赏南宫墨哈谯笪年爱阳佟第五言福"
)

NON_NAME_WORDS = set(
    "甲方 乙方 丙方 丁方 以上 以下 因为 所以 虽然 如果 可以 应当 必须 需要 能够 可能 应该 "
    "没有 不是 这个 那个 这些 那些 什么 怎么 如何 关于 对于 根据 按照 依照 通过 本次 "
    "该案 本案 合同 协议 条款 约定 规定 办法 制度 规则 文件 通知 报告 申请 批复 决定 意见 "
    "函件 公告 公示 第一 第二 第三 第四 第五 第六 第七 第八 第九 第十 首先 其次 再次 最后 "
    "合计 共计 总计 小计 金额 数量 日期 时间 地点 方式 标准 条件 范围 内容 目的 意义 权利 "
    "义务 责任 风险 费用 损失 赔偿 违约金 双方 各方 单方 对方 本人 我方 贵方 年 月 日 时 分 "
    "秒 元 角 分 整 买方 卖方 出租方 承租方 转让方 受让方 发包人 承包人 委托人 受托人 "
    "保证人 抵押人 出质人 债务人 债权人 法定代表人 委托代理人 授权代表 联系人 见证人 "
    "担保人 当事人 第三人 主体 名称 住所 地址 电话 邮箱 账号 银行 "
    "响应人 响应方 响应 采购人 采购 招标人 招标 投标人 投标 中标人 中标 "
    "成交人 成交 供应商 谈判 竞争 评审 竞争性 磋商 询价 "
    "公告 邀请 邀请函 通知 目录 总则 附件 附表 附录 "
    "总价 单价 合价 金额 合计 总计 共计 小计 累计 总额 "
    "人民币 美元 欧元 英镑 日元 港币 "
    "含税 不含税 税前 税后 税率 税额 税种 税收 税务 增值税 所得税 "
    "百分之 千分之 万分之 百分比 百分点 比率 比例 系数 "
    "年利率 月利率 日利率 利率 利息 本息 本金 "
    "不可抗力 违约金 赔偿金 保证金 定金 押金 质保金 履约金 "
    "验收 检验 检测 试验 调试 试运行 交付 交货 收货 "
    "保密 知识产权 专有技术 技术秘密 商业秘密 技术资料 "
    "仲裁 诉讼 调解 和解 协商 谈判 管辖 法律适用 "
    "生效 失效 终止 解除 变更 修改 补充 续签 展期 "
    "通知 送达 告知 确认 同意 认可 批准 核准 审批 备案 "
    "签字 盖章 签章 签署 签收 签认 签批 签发 "
    "注册资本 实收资本 总资产 净资产 营业收入 净利润 "
    "正本 副本 原件 复印件 扫描件 电子版 纸质版 "
    "国家标准 行业标准 企业标准 地方标准 团体标准 国际标准 "
    "技术规范 技术条件 技术参数 技术要求 技术指标 技术标准 "
    "规格 型号 品牌 产地 制造商 生产商 供应商 经销商 代理商 "
    "编制 制定 起草 拟定 撰写 出具 提供 提交 递交 "
    "审核 审查 审批 批准 核准 备案 登记 注册 "
    "分析 评估 评价 判断 决定 意见 建议 方案 "
    "加密 解密 密码 密钥 数字证书 电子签名 "
    "确定 确认 选定 选择 推荐 排序 排名 "
    "签约 签订 签署 签章 盖章 "
    "采购内容 采购范围 采购清单 采购数量 采购规格 "
    "合同资料表 合同条款 合同条件 合同格式 合同模板 "
    "报价表 报价单 报价函 报价信 报价书 "
    "响应函 响应书 响应文件 响应方案 响应计划 "
    "资格审查 资格预审 资格后审 资格条件 资格要求 "
    "符合性 响应性 实质性 形式性 资格性 技术性 商务性 "
    "偏差 偏离 修改 调整 变更 澄清 说明 补正 "
    "撤回 撤销 取消 放弃 拒绝 接受 同意 "
    "澄清 说明 补正 修正 更改 替换 "
    "现场 实地 书面 口头 电子 线上 线下 网上 网下 "
    "密封 开启 公开 公示 公告 公布 发布 "
    "监督管理 监督检查 监督 管理 审查 审核 复核 "
    "投诉 质疑 异议 申诉 举报 反映 "
    "暂停 中止 终止 停止 继续 恢复 重启 "
    "有效 无效 生效 失效 过期 逾期 到期 届满 "
    "工作日 自然日 日历日 营业日 节假日 休息日 "
    "前 后 内 外 上 下 左 右 中 间 旁 侧 边 "
    "一份 二份 三份 多份 一式 一致 正本 副本 复印件 "
    "中文 英文 中英文 双语 对照 翻译 译本 "
    "具有 享有 拥有 持有 获得 取得 得到 "
    "按照 依照 根据 依据 参照 参考 比照 对照 "
    "包括 包含 涵盖 涉及 限于 除非 除 外 "
    "以及 和 与 或 及 并 且 而 但 然而 "
    "如果 若 如 则 但 然而 且 或 以及 并 "
    "上述 以下 如下 如上 如前 其后 此前 此后 届时 当即 "
    "本文件 本协议 本合同 本附件 本条款 本约定 本条件 "
    "其他 其余 剩余 额外 附加 补充 附属 辅助 "
    "相关 关联 对应 相应 匹配 适配 适合 适用 "
    "具体 明确 清晰 详细 详尽 完整 全面 系统 "
    "必要 充分 足够 适当 合理 合规 合法 有效 "
    "最大 最小 最多 最少 最高 最低 最早 最晚 "
    "首次 再次 多次 每次 单次 累计 连续 持续 "
    "重新 再次 重复 反复 多次 另行 单独 分别 "
    "共同 一致 统一 联合 协同 合作 配合 协调 "
    "自动 手动 人工 智能 半自动 全自动 机械 电子 "
    "书面 口头 电子 线上 线下 网络 现场 远程 "
    "致 各 各供应商 各响应人 "
    "保证金 质保金 押金 定金 履约保证金 投标保证金 "
    "偏离表 差异表 逐项 逐条 对照 比较 "
    "补充材料 补充文件 补充说明 补充资料 "
    "资格证明 资格文件 资质证书 营业执照 许可证 "
    "书 函 函件 信 信函 表 单 证 文件 资料 "
    "目录 索引 标题 章节 页 页码 编号 "
    "见 详见 参见 参照 按照 依照 根据 "
    "报价一览表 分项报价表 报价明细表 报价汇总表 "
    "技术偏离表 商务偏离表 合同资料表 合同条款 "
    "主要合同条款 采购内容 采购范围 采购清单 采购数量 采购规格 "
    "资格证明文件 资格文件 资质文件 证明文件 "
    "其他补充材料 其他材料 补充材料 补充文件 "
    "响应文件 响应函 响应书 响应方案 响应计划 "
    "竞争谈判 竞争性谈判 竞争性磋商 竞争谈判采购 "
    "确定成交 成交人 成交供应商 成交结果 "
    "响应文件编制 响应文件加密 响应文件递交 "
    "及构成 构成 组成 包含 包括 "
    "第 章 节 条 款 项 项目 附件 附表 附录 "
    .split()
)


class ReplaceItem:
    def __init__(self, original, category, enabled=True):
        self.original = original
        self.category = category
        self.enabled = enabled


class ReplaceListDialog(QDialog):
    def __init__(self, replace_items, parent=None):
        super().__init__(parent)
        self.replace_items = replace_items
        self.setWindowTitle("替换清单预览")
        self.resize(900, 500)
        self.setModal(True)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        info = QLabel("以下为识别到的敏感实体，可手动调整后点击[重新生成]：")
        info.setWordWrap(True)
        layout.addWidget(info)

        self.table = QTableWidget(len(self.replace_items), 4)
        self.table.setHorizontalHeaderLabels(["启用", "原文", "类别", "替换为"])
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.table.setEditTriggers(QTableWidget.DoubleClicked)

        for i, item in enumerate(self.replace_items):
            cb = QCheckBox()
            cb.setChecked(item.enabled)
            self.table.setCellWidget(i, 0, cb)
            text_item = QTableWidgetItem(item.original)
            text_item.setFlags(text_item.flags() | Qt.ItemIsEditable)
            self.table.setItem(i, 1, text_item)
            self.table.setItem(i, 2, QTableWidgetItem(item.category))
            mask_item = QTableWidgetItem(MASK)
            mask_item.setFlags(mask_item.flags() & ~Qt.ItemIsEditable)
            self.table.setItem(i, 3, mask_item)

        layout.addWidget(self.table)
        btn_layout = QHBoxLayout()
        self.btn_select_all = QPushButton("全选")
        self.btn_unselect_all = QPushButton("全不选")
        self.btn_regenerate = QPushButton("重新生成")
        self.btn_cancel = QPushButton("取消")
        self.btn_select_all.clicked.connect(self._select_all)
        self.btn_unselect_all.clicked.connect(self._unselect_all)
        self.btn_regenerate.clicked.connect(self.accept)
        self.btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(self.btn_select_all)
        btn_layout.addWidget(self.btn_unselect_all)
        btn_layout.addStretch()
        btn_layout.addWidget(self.btn_regenerate)
        btn_layout.addWidget(self.btn_cancel)
        layout.addLayout(btn_layout)

    def _select_all(self):
        for i in range(self.table.rowCount()):
            w = self.table.cellWidget(i, 0)
            if w and isinstance(w, QCheckBox):
                w.setChecked(True)

    def _unselect_all(self):
        for i in range(self.table.rowCount()):
            w = self.table.cellWidget(i, 0)
            if w and isinstance(w, QCheckBox):
                w.setChecked(False)

    def get_updated_items(self):
        items = []
        for i in range(self.table.rowCount()):
            cb = self.table.cellWidget(i, 0)
            enabled = cb.isChecked() if cb else True
            original = self.table.item(i, 1).text() if self.table.item(i, 1) else ""
            category = self.table.item(i, 2).text() if self.table.item(i, 2) else ""
            items.append(ReplaceItem(original, category, enabled))
        return items


class Desensitizer:
    def __init__(self, custom_dict=None, enabled_types=None, sensitivity="标准"):
        self.mapping = MappingManager()
        self.custom_dict = [s.strip() for s in (custom_dict or []) if s.strip()]
        self.enabled_types = enabled_types or {}
        self.sensitivity = sensitivity
        self._value_to_ph = {}
        self._value_to_type = {}
        self._count = 0
        self._alerts = []
        self._readability_issues = []
        self._replace_items = []
        self._compile_patterns()

    def _compile_patterns(self):
        self._phone_re = re.compile(r"1[3-9]\d{9}")
        self._id_re = re.compile(r"\d{17}[\dXx]")
        self._email_re = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
        self._card_re = re.compile(r"(?<!\d)\d{16,19}(?!\d)")
        self._tel_re = re.compile(r"0\d{2,3}-\d{7,8}")
        self._usid_re = re.compile(r"\d{3}-\d{2}-\d{4}")
        self._postal_re = re.compile(r"(?<!\d)[1-9]\d{5}(?!\d)")
        self._url_re = re.compile(r"https?://[^\s\u4e00-\u9fff，。、；：()]+")
        self._www_re = re.compile(r"(?:www\.)[a-zA-Z0-9][a-zA-Z0-9.-]*\.[a-zA-Z]{2,}(?:/[^\s\u4e00-\u9fff，。、；：()]*)?")
        self._company_re = re.compile(r"[\u4e00-\u9fff]{2,}(?:" + COMPANY_SUFFIX + ")")
        self._company_en_re = re.compile(r"[A-Z][a-zA-Z\s&]+(?:Inc|Corp|LLC|Ltd|PLC|GmbH|AG|SA|SAS|SARL|BV|NV|KK|Co|Company|Limited|Incorporated|Group|Holdings|Partners|Associates|Consulting|Management|Services|Capital|Investments|Technologies|Solutions|Systems|Industries|Enterprises|Ventures|Institute|University|College|School|Academy|Foundation|Laboratories|Laboratory)")
        self._swift_re = re.compile(r"\b[A-Z]{6}[A-Z0-9]{2}(?:[A-Z0-9]{3})?\b")
        self._iban_re = re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{1,30}\b")
        self._address_chain_re = re.compile(r"(?:[\u4e00-\u9fff]{2,6}(?:省|自治区)\s*)?[\u4e00-\u9fff]{2,7}(?:市|自治州|区)\s*[\u4e00-\u9fff]{2,7}(?:区|县|市|镇|乡|街道|新城|开发区|高新区|工业园|产业园|科技园|新区|经开区)(?:\s*[\u4e00-\u9fff]{2,8}(?:路|街|道|巷|大道|大街|工业园|开发区|产业园|科技园|区|村|镇|乡|号|大道|公路|线|口岸|港|枢纽|center)){0,3}")
        self._star_re = re.compile(r"\*+")

    def is_whitelisted(self, text):
        if text in WHITELIST_WORDS:
            return True
        for pattern in WHITELIST_PATTERNS:
            m = pattern.fullmatch(text)
            if m:
                return True
            m = pattern.search(text)
            if m and m.group(0) == text:
                return True
        return False

    def is_common_term(self, text):
        return text in COMMON_TERMS or text in COMMON_TERMS_EN

    def build_mapping(self, full_text):
        self._value_to_ph.clear()
        self._value_to_type.clear()
        self._count = 0
        self._alerts = []
        self._readability_issues = []
        self._replace_items = []
        self.mapping.clear()
        if self.sensitivity == "保守":
            allowed = {"人名", "身份证", "电话", "银行账号", "网址", "自定义词条"}
            self.enabled_types = {k: k in allowed for k in self.enabled_types}
        elif self.sensitivity == "激进":
            self.enabled_types = {k: True for k in self.enabled_types}
        self._apply_custom_dict(full_text)
        self._extract_field_values(full_text)
        self._apply_regex_fallback(full_text)
        self._extract_address_chain(full_text)
        self._extract_qq_meeting(full_text)
        self._filter_common_terms()
        self._build_replace_items()
        return dict(self._value_to_ph)

    def desensitize(self, text, enabled_overrides=None):
        if not self._value_to_ph:
            return text
        if enabled_overrides:
            overrides = {item.original: item.enabled for item in enabled_overrides}
            for val in sorted(self._value_to_ph.keys(), key=len, reverse=True):
                if val in text:
                    enabled = overrides.get(val, True)
                    if enabled:
                        text = text.replace(val, MASK)
        else:
            for val in sorted(self._value_to_ph.keys(), key=len, reverse=True):
                if val in text:
                    text = text.replace(val, MASK)
        return text

    def get_mapping_summary(self):
        return "脱敏处理完成，共替换 %d 处敏感实体" % self._count

    def get_alerts(self):
        return list(self._alerts)

    def get_readability_issues(self):
        return list(self._readability_issues)

    def get_replace_items(self):
        return list(self._replace_items)

    def _filter_common_terms(self):
        all_terms = set(COMMON_TERMS) | set(COMMON_TERMS_EN)
        to_remove = []
        for value in list(self._value_to_ph.keys()):
            if value in all_terms or self.is_common_term(value):
                to_remove.append(value)
        for value in to_remove:
            del self._value_to_ph[value]
            self._value_to_type.pop(value, None)
            if value in self.mapping._map:
                del self.mapping._map[value]
            self._count -= 1
            self._readability_issues.append("已自动还原常见合同术语：" + value)

    def _build_replace_items(self):
        for value, ptype in sorted(self._value_to_type.items(), key=lambda x: -len(x[0])):
            self._replace_items.append(ReplaceItem(value, ptype, True))

    def _apply_custom_dict(self, full_text):
        for entry in self.custom_dict:
            if not entry or entry in self._value_to_ph or entry not in full_text:
                continue
            if self.is_whitelisted(entry) or self.is_common_term(entry):
                continue
            self._add_entity(entry, "自定义词条")

    def _extract_field_values(self, full_text):
        label_alt = "|".join(re.escape(l) for l in BLACKLIST_LABELS_SORTED)
        pat = re.compile(r"(?P<label>" + label_alt + r")\s*[:：\-—|]\s*(?P<value>[^" + NL + r"]*?)(?=\s*(?:" + label_alt + r")\s*[:：\-—]|" + NL + r"|$)")
        for m in pat.finditer(full_text):
            label = m.group("label")
            raw = m.group("value").strip()
            if not raw:
                continue
            value = self._truncate_at_next_label(raw)
            ptype = self._map_label_to_type(BLACKLIST_LABELS.get(label, "未知"))
            self._add_value(value, ptype)
        pat2 = re.compile(r"(?m)^\s*(?P<label>" + label_alt + r")\s{1,}(?P<value>[^" + NL + r"]*)$")
        for m in pat2.finditer(full_text):
            label = m.group("label")
            raw = m.group("value").strip()
            if not raw:
                continue
            value = self._truncate_at_next_label(raw)
            ptype = self._map_label_to_type(BLACKLIST_LABELS.get(label, "未知"))
            if self._value_looks_entity(value, ptype):
                self._add_value(value, ptype)
        lines = full_text.split(NL)
        for i, line in enumerate(lines):
            ls = line.strip().rstrip("：: 	")
            if not ls or ls not in BLACKLIST_LABELS:
                continue
            if i + 1 >= len(lines):
                continue
            nxt = lines[i + 1].strip()
            if not nxt or nxt.strip().rstrip("：: 	") in BLACKLIST_LABELS:
                continue
            value = self._truncate_at_next_label(nxt)
            ptype = self._map_label_to_type(BLACKLIST_LABELS.get(ls, "未知"))
            self._add_value(value, ptype)

    def _map_label_to_type(self, raw_type):
        mapping = {
            "人名": "人名", "电话": "电话", "邮箱": "邮箱",
            "地址": "地址", "网址": "网址", "身份证": "身份证",
            "信用代码": "信用代码", "证照号": "信用代码",
            "银行账号": "银行账号", "银行卡": "银行账号", "开户行": "银行账号",
            "SWIFT": "SWIFT", "IBAN": "SWIFT",
            "合同编号": "合同编号", "案号": "合同编号",
            "订单号": "合同编号", "发票号": "合同编号",
            "项目": "合同编号", "编号": "合同编号", "邮编": "邮编",
        }
        return mapping.get(raw_type, raw_type)

    def _truncate_at_next_label(self, value):
        best = None
        for label in BLACKLIST_LABELS_SORTED:
            idx = value.find(label)
            if idx > 0 and (best is None or idx < best):
                best = idx
        if best is not None:
            return value[:best].strip()
        return value

    def _value_looks_entity(self, value, ptype):
        if ptype in ("银行账号",):
            return bool(re.match(r"^[\d\s\-]{6,}$", value))
        if ptype == "人名":
            return bool(re.match(r"^[\u4e00-\u9fff]{2,6}$", value)) and value[0] in COMMON_SURNAMES
        if ptype == "公司名":
            return len(value) >= 4
        if ptype == "电话":
            return bool(re.match(r"^[\d+\-()\s]{6,}$", value))
        if ptype == "邮箱":
            return "@" in value
        if ptype == "地址":
            return len(value) >= 6
        if ptype == "网址":
            return "http" in value or "www." in value or "." in value
        return len(value) >= 2

    def _add_value(self, value, ptype):
        value = value.strip().rstrip("（(【[,，、；;:：。")
        if not value or len(value) > 100:
            return
        if self.is_whitelisted(value) or self.is_common_term(value):
            return
        numeric_types = ("银行账号", "电话", "身份证", "信用代码", "邮编", "合同编号", "SWIFT")
        if ptype in numeric_types:
            self._add_entity(value, ptype)
            return
        if re.match(r"^[\d\s,\u3000、；;\:\-—~～]+$", value):
            return
        if len(value) <= 1:
            return
        self._add_entity(value, ptype)

    def _extract_qq_meeting(self, full_text):
        pat = re.compile(r"(?:QQ群|QQ群号|QQ群号码|腾讯会议|腾讯会议号|腾讯会议号码|会议号|会议ID|Meeting ID|Meeting No|Meeting Number)\s*[:：]?\s*(\d{5,11})")
        for m in pat.finditer(full_text):
            digit = m.group(1).strip()
            if digit and digit not in self._value_to_ph:
                self._add_entity(digit, "编号")

    def _extract_address_chain(self, full_text):
        for m in self._address_chain_re.finditer(full_text):
            full_name = m.group(0).strip()
            if len(full_name) <= 6 or full_name in self._value_to_ph:
                continue
            if self.is_whitelisted(full_name) or self.is_common_term(full_name):
                continue
            self._add_entity(full_name, "地址")

    def _apply_regex_fallback(self, full_text):
        for m in self._url_re.finditer(full_text):
            v = m.group(0).strip().rstrip(",.;:!?")
            if v not in self._value_to_ph and len(v) > 8:
                if self.is_whitelisted(v) or self.is_common_term(v):
                    continue
                self._add_entity(v, "网址")
        for m in self._www_re.finditer(full_text):
            v = m.group(0).strip().rstrip(",.;:!?")
            if v not in self._value_to_ph and len(v) > 5:
                if self.is_whitelisted(v) or self.is_common_term(v):
                    continue
                self._add_entity(v, "网址")
        for m in self._phone_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                self._add_entity(v, "电话")
        for m in self._id_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                self._add_entity(v, "身份证")
        for m in self._email_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                self._add_entity(v, "邮箱")
        for m in self._card_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                self._add_entity(v, "银行账号")
        for m in self._tel_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                self._add_entity(v, "电话")
        for m in self._usid_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                self._add_entity(v, "身份证")
        for m in self._postal_re.finditer(full_text):
            v = m.group(0)
            if v not in self._value_to_ph:
                if self.is_whitelisted(v) or self.is_common_term(v):
                    continue
                self._add_entity(v, "邮编")
        for m in self._swift_re.finditer(full_text):
            v = m.group(0)
            if v in self._value_to_ph:
                continue
            if self.is_whitelisted(v) or self.is_common_term(v):
                continue
            self._add_entity(v, "SWIFT")
        for m in self._iban_re.finditer(full_text):
            v = m.group(0)
            if v in self._value_to_ph:
                continue
            if self.is_whitelisted(v) or self.is_common_term(v):
                continue
            if len(v) < 8:
                continue
            self._add_entity(v, "SWIFT")
        for m in self._company_re.finditer(full_text):
            v = m.group(0)
            if v in self._value_to_ph:
                continue
            if v in COMPANY_EXCLUDE_SET:
                continue
            if len(v) < 4:
                continue
            if self.is_whitelisted(v) or self.is_common_term(v):
                continue
            self._add_entity(v, "公司名")
        for m in self._company_en_re.finditer(full_text):
            v = m.group(0).strip()
            if v in self._value_to_ph:
                continue
            if len(v) < 5:
                continue
            if self.is_whitelisted(v) or self.is_common_term(v):
                continue
            self._add_entity(v, "公司名")

    def _add_entity(self, value, ptype):
        if value in self._value_to_ph:
            return
        self._value_to_ph[value] = MASK
        self._value_to_type[value] = ptype
        self.mapping.add(value, MASK)
        self._count += 1

    def run_validation(self, original_text, desensitized_text):
        alerts = []
        if self._star_re.search(desensitized_text):
            alerts.append("【脱敏异常告警，请人工通读复核文档后再流转】检测到星号字符")
        for keyword in ["民法典", "合同法", "不可抗力", "违约金", "仲裁", "响应人", "谈判", "评审", "成交人", "采购"]:
            if keyword in original_text and keyword not in desensitized_text:
                alerts.append("【可读性提示】白名单关键词被误替换: " + keyword)
                break
        for issue in self._readability_issues:
            alerts.append("【可读性复检】" + issue)
        return alerts


class Restorer:
    def __init__(self):
        self.mapping = MappingManager()

    def load_mapping(self, path):
        self.mapping = MappingManager.load_from_file(path)
        return self.mapping.count > 0

    def restore(self, text):
        rev = self.mapping.get_all_reverse()
        for ph in sorted(rev.keys(), key=len, reverse=True):
            text = text.replace(ph, rev[ph])
        return text

    def find_mapping_file(self, desensitized_file):
        d = os.path.dirname(desensitized_file)
        b = os.path.basename(desensitized_file)
        candidates = []
        if b.startswith("脱敏_"):
            rest = b[3:]
            name = os.path.splitext(rest)[0]
            candidates.append(os.path.join(d, "mapping_%s.json" % name))
        candidates.append(os.path.join(d, "mapping.json"))
        for c in candidates:
            if os.path.exists(c):
                return c
        return None


class DocParser:
    def __init__(self):
        self.paragraphs = []
        self._original_doc = None

    def load(self, path):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx":
            return self._load_docx(path)
        elif ext == ".xlsx":
            return self._load_xlsx(path)
        elif ext == ".pdf":
            return self._load_pdf(path)
        else:
            return self._load_txt(path)

    def _load_xlsx(self, path):
        try:
            import openpyxl
        except ImportError:
            return False
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        self.paragraphs = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() if c else "" for c in row]
                t = " | ".join(c for c in cells if c)
                if t.strip():
                    self.paragraphs.append({"text": t, "type": "xlsx_row"})
        wb.close()
        return len(self.paragraphs) > 0

    def _load_docx(self, path):
        from docx import Document
        self._original_doc = Document(path)
        self.paragraphs = []
        for para in self._original_doc.paragraphs:
            if para.text.strip():
                self.paragraphs.append({"text": para.text, "type": "paragraph", "_para": para})
        for table in self._original_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        self.paragraphs.append({"text": cell.text.strip(), "type": "cell", "_cell": cell})
        try:
            for section in self._original_doc.sections:
                header = section.header
                if header and header.paragraphs:
                    for para in header.paragraphs:
                        if para.text.strip():
                            self.paragraphs.append({"text": para.text.strip(), "type": "header", "_para": para})
                footer = section.footer
                if footer and footer.paragraphs:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            self.paragraphs.append({"text": para.text.strip(), "type": "footer", "_para": para})
        except Exception:
            pass
        return len(self.paragraphs) > 0

    def _load_pdf(self, path):
        try:
            import pdfplumber
        except ImportError:
            return False
        self.paragraphs = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for para in text.split(NL + NL):
                        para = para.strip()
                        if para:
                            self.paragraphs.append({"text": para, "type": "paragraph"})
        return len(self.paragraphs) > 0

    def _load_txt(self, path):
        self.paragraphs = []
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        for para in content.split(NL):
            para = para.strip()
            if para:
                self.paragraphs.append({"text": para, "type": "paragraph"})
        return len(self.paragraphs) > 0

    def get_text(self):
        return NL.join(p["text"] for p in self.paragraphs)

    def get_paragraphs(self):
        return list(self.paragraphs)

    def save_docx(self, path, texts):
        if self._original_doc is None:
            from docx import Document
            doc = Document()
            for text in texts:
                doc.add_paragraph(text)
            doc.save(path)
            return
        for para_info, new_text in zip(self.paragraphs, texts):
            if "_para" in para_info:
                para = para_info["_para"]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)
            elif "_cell" in para_info:
                cell = para_info["_cell"]
                cell.text = new_text
        self._original_doc.save(path)

    def save_txt(self, path, texts):
        with open(path, "w", encoding="utf-8") as f:
            for text in texts:
                f.write(text + NL)


# ============================================================
# 拖拽表格控件（修复拖拽功能）
# ============================================================
class DropTableWidget(QTableWidget):
    """支持文件拖拽的表格控件"""
    file_dropped = Signal(list)  # 信号：拖入的文件路径列表

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setAcceptDrops(True)
        self.setDragDropMode(QAbstractItemView.DropOnly)
        self.setDefaultDropAction(Qt.CopyAction)
        self._drag_over = False

    def dragEnterEvent(self, event):
        """鼠标拖入时触发"""
        if event.mimeData().hasUrls():
            # 检查是否有文件URL
            has_files = any(url.isLocalFile() for url in event.mimeData().urls())
            if has_files:
                event.acceptProposedAction()
                self._drag_over = True
                self.setStyleSheet("QTableWidget { border: 2px dashed #4A90D9; background: #E8F4FD; }")
                return
        event.ignore()
        super().dragEnterEvent(event)

    def dragMoveEvent(self, event):
        """鼠标在控件上移动时持续触发（保持接受状态）"""
        if event.mimeData().hasUrls():
            has_files = any(url.isLocalFile() for url in event.mimeData().urls())
            if has_files:
                event.acceptProposedAction()
                return
        event.ignore()

    def dragLeaveEvent(self, event):
        """鼠标离开控件时恢复样式"""
        self._drag_over = False
        self.setStyleSheet("")
        super().dragLeaveEvent(event)

    def dropEvent(self, event):
        """释放鼠标时处理文件"""
        self._drag_over = False
        self.setStyleSheet("")  # 恢复正常样式
        files = []
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if url.isLocalFile():
                    path = url.toLocalFile()
                    if os.path.isfile(path):
                        files.append(path)
        if files:
            event.acceptProposedAction()
            self.file_dropped.emit(files)
        else:
            event.ignore()
        super().dropEvent(event)


class LogSignal(QObject):
    append = Signal(str)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.config = AppConfig()
        self.setWindowTitle(APP_NAME + " " + APP_VERSION)
        self.resize(1200, 800)
        self.custom_dict = []
        self.log_signal = LogSignal()
        self.log_signal.append.connect(self._append_log)
        self._processing = False
        self.init_ui()
        self._setup_shortcuts()

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("Ctrl+O"), self, self.on_add_files)
        QShortcut(QKeySequence("Ctrl+P"), self, self.on_start)
        QShortcut(QKeySequence("Ctrl+E"), self, self._open_output_dir)
        QShortcut(QKeySequence("Delete"), self, self._remove_selected_file)

    # ===== 拖拽功能（已迁移至 DropTableWidget 子类）=====
    # 旧方法 _setup_drag_drop / _drag_enter / _drop 已移除
    # 现在使用 DropTableWidget.file_dropped 信号 → _on_files_dropped

    def _on_files_dropped(self, file_paths):
        """处理拖拽放入的文件列表（由 DropTableWidget 信号触发）"""
        added = 0
        for path in file_paths:
            ext = os.path.splitext(path)[1].lower()
            if ext in ('.docx', '.doc', '.pdf', '.xlsx', '.txt'):
                self._add_single_file(path)
                added += 1
            else:
                self.log("跳过不支持的文件格式：%s" % os.path.basename(path))
        if added > 0:
            self.log("拖拽添加了 %d 个文件" % added)

    def init_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setSpacing(8)
        root.setContentsMargins(8, 8, 8, 8)
        main_splitter = QSplitter(Qt.Horizontal)
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(8)

        top_row = QHBoxLayout()
        mode_box = QGroupBox("任务模式")
        mode_layout = QHBoxLayout(mode_box)
        self.radio_de = QRadioButton("脱敏")
        self.radio_re = QRadioButton("还原")
        self.radio_de.setChecked(True)
        self.mode_group = QButtonGroup(self)
        self.mode_group.addButton(self.radio_de)
        self.mode_group.addButton(self.radio_re)
        mode_layout.addWidget(self.radio_de)
        mode_layout.addWidget(self.radio_re)
        mode_layout.addStretch()
        top_row.addWidget(mode_box)

        sens_box = QGroupBox("敏感度")
        sens_layout = QHBoxLayout(sens_box)
        self.sens_combo = QComboBox()
        self.sens_combo.addItems(["保守", "标准", "激进"])
        self.sens_combo.setCurrentText(self.config.sensitivity)
        self.sens_combo.currentTextChanged.connect(self._on_sens_change)
        sens_layout.addWidget(self.sens_combo)
        sens_layout.addStretch()
        top_row.addWidget(sens_box)
        left_layout.addLayout(top_row)

        file_box = QGroupBox("待处理文件（支持拖拽添加）")
        file_layout = QVBoxLayout(file_box)
        file_layout.setSpacing(6)
        ops = QHBoxLayout()
        self.btn_add = QPushButton("添加文档")
        self.btn_add.setObjectName("btn_add")
        self.btn_clear = QPushButton("清空列表")
        self.btn_clear.setObjectName("btn_clear")
        self.btn_add.clicked.connect(self.on_add_files)
        self.btn_clear.clicked.connect(self.on_clear_list)
        ops.addWidget(self.btn_add)
        ops.addWidget(self.btn_clear)
        ops.addStretch()
        file_layout.addLayout(ops)
        # 使用支持拖拽的自定义表格控件（修复拖拽功能）
        self.table = DropTableWidget(0, 5)
        self.table.setHorizontalHeaderLabels(["文件名", "类型", "大小", "状态", "路径"])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self._show_file_context_menu)
        # 连接拖拽信号到文件添加方法
        self.table.file_dropped.connect(self._on_files_dropped)
        file_layout.addWidget(self.table, stretch=1)

        out_row = QHBoxLayout()
        out_row.addWidget(QLabel("输出目录："))
        self.output_dir = QLineEdit()
        self.output_dir.setPlaceholderText("默认在源文件旁创建任务文件夹")
        self.output_dir.setText(self.config.last_output_dir)
        self.btn_choose_dir = QPushButton("选择目录")
        self.btn_choose_dir.setObjectName("btn_choose_dir")
        self.btn_choose_dir.clicked.connect(self.on_choose_dir)
        self.btn_open_output = QPushButton("打开文件夹")
        self.btn_open_output.setObjectName("btn_open_output")
        self.btn_open_output.clicked.connect(self._open_output_dir)
        out_row.addWidget(self.output_dir, stretch=1)
        out_row.addWidget(self.btn_choose_dir)
        out_row.addWidget(self.btn_open_output)
        file_layout.addLayout(out_row)
        left_layout.addWidget(file_box, stretch=1)

        action_layout = QHBoxLayout()
        self.btn_process = QPushButton("开始处理")
        self.btn_process.setObjectName("btn_process")
        self.btn_process.setMinimumWidth(160)
        self.btn_process.clicked.connect(self.on_start)
        self.progress = QProgressBar()
        self.progress.setValue(0)
        self.progress.setRange(0, 100)
        action_layout.addWidget(self.btn_process)
        action_layout.addWidget(self.progress, stretch=1)

        # ===== 补丁 #3：关于按钮 & 菜单栏 =====
        self.btn_about = QPushButton("关于")
        self.btn_about.setObjectName("btn_choose_dir")
        self.btn_about.clicked.connect(self.open_about)
        action_layout.addWidget(self.btn_about)
        left_layout.addLayout(action_layout)

        # 创建菜单栏
        menubar = self.menuBar()
        help_menu = menubar.addMenu("帮助(H)")
        act_about = QAction("关于 " + APP_NAME, self)
        act_about.triggered.connect(self.open_about)
        help_menu.addAction(act_about)
        act_update = QAction("检查更新", self)
        act_update.triggered.connect(self.check_update)
        help_menu.addAction(act_update)
        help_menu.addSeparator()
        act_cfg = QAction("打开配置目录", self)
        act_cfg.triggered.connect(self._open_config_dir)
        help_menu.addAction(act_cfg)

        self.log_panel = QPlainTextEdit()
        self.log_panel.setReadOnly(True)
        self.log_panel.setMaximumBlockCount(2000)
        self.log_panel.setPlaceholderText("处理日志将在此显示……")
        left_layout.addWidget(self.log_panel, stretch=1)

        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_tabs = QTabWidget()

        type_tab = QWidget()
        type_layout = QVBoxLayout(type_tab)
        type_group = QGroupBox("脱敏类型控制（勾选=脱敏，取消=保留）")
        type_grid = QGridLayout(type_group)
        self.type_checks = {}
        type_names = [
            ("人名", 0, 0), ("公司名", 0, 1), ("电话", 0, 2),
            ("邮箱", 1, 0), ("地址", 1, 1), ("身份证", 1, 2),
            ("银行账号", 2, 0), ("信用代码", 2, 1), ("合同编号", 2, 2),
            ("邮编", 3, 0), ("SWIFT", 3, 1), ("网址", 3, 2),
            ("自定义词条", 4, 0),
        ]
        for name, r, c in type_names:
            cb = QCheckBox(name)
            cb.setChecked(self.config.enabled_types.get(name, True))
            cb.stateChanged.connect(self._on_type_check_changed)
            self.type_checks[name] = cb
            type_grid.addWidget(cb, r, c)
        type_layout.addWidget(type_group)
        type_layout.addStretch()
        right_tabs.addTab(type_tab, "脱敏类型")

        dict_tab = QWidget()
        dict_layout = QVBoxLayout(dict_tab)
        dict_layout.setSpacing(6)

        wl_box = QGroupBox("白名单（强制不脱敏）")
        wl_layout = QVBoxLayout(wl_box)
        self.wl_list = QListWidget()
        self.wl_list.setAlternatingRowColors(True)
        for w in self.config.custom_whitelist:
            self.wl_list.addItem(QListWidgetItem(w))
        wl_layout.addWidget(self.wl_list, stretch=1)
        wl_input_row = QHBoxLayout()
        self.wl_input = QLineEdit()
        self.wl_input.setPlaceholderText("输入白名单词条，回车添加...")
        self.wl_input.returnPressed.connect(self._add_wl_from_input)
        self.btn_wl_add = QPushButton("+")
        self.btn_wl_add.setObjectName("btn_wl_add")
        self.btn_wl_add.clicked.connect(self._add_wl_from_input)
        self.btn_wl_del = QPushButton("-")
        self.btn_wl_del.setObjectName("btn_wl_del")
        self.btn_wl_del.clicked.connect(self._del_wl_selected)
        wl_input_row.addWidget(self.wl_input, stretch=1)
        wl_input_row.addWidget(self.btn_wl_add)
        wl_input_row.addWidget(self.btn_wl_del)
        wl_layout.addLayout(wl_input_row)
        dict_layout.addWidget(wl_box, stretch=1)

        bl_box = QGroupBox("黑名单（强制脱敏）")
        bl_layout = QVBoxLayout(bl_box)
        self.bl_list = QListWidget()
        self.bl_list.setAlternatingRowColors(True)
        for b in self.config.custom_blacklist:
            self.bl_list.addItem(QListWidgetItem(b))
        bl_layout.addWidget(self.bl_list, stretch=1)
        bl_input_row = QHBoxLayout()
        self.bl_input = QLineEdit()
        self.bl_input.setPlaceholderText("输入黑名单词条，回车添加...")
        self.bl_input.returnPressed.connect(self._add_bl_from_input)
        self.btn_bl_add = QPushButton("+")
        self.btn_bl_add.setObjectName("btn_bl_add")
        self.btn_bl_add.clicked.connect(self._add_bl_from_input)
        self.btn_bl_del = QPushButton("-")
        self.btn_bl_del.setObjectName("btn_bl_del")
        self.btn_bl_del.clicked.connect(self._del_bl_selected)
        bl_input_row.addWidget(self.bl_input, stretch=1)
        bl_input_row.addWidget(self.btn_bl_add)
        bl_input_row.addWidget(self.btn_bl_del)
        bl_layout.addLayout(bl_input_row)
        dict_layout.addWidget(bl_box, stretch=1)

        btn_row = QHBoxLayout()
        self.btn_load_dict = QPushButton("从文件加载")
        self.btn_load_dict.setObjectName("btn_choose_dir")
        self.btn_export_dict = QPushButton("导出词条")
        self.btn_export_dict.setObjectName("btn_choose_dir")
        self.btn_apply_dict = QPushButton("应用到本次任务")
        self.btn_load_dict.clicked.connect(self.on_load_dict_file)
        self.btn_export_dict.clicked.connect(self._export_custom_dict)
        self.btn_apply_dict.clicked.connect(self.on_apply_custom_dict)
        btn_row.addWidget(self.btn_load_dict)
        btn_row.addWidget(self.btn_export_dict)
        btn_row.addWidget(self.btn_apply_dict)
        dict_layout.addLayout(btn_row)
        right_tabs.addTab(dict_tab, "自定义词条")

        biz_tab = QWidget()
        biz_layout = QVBoxLayout(biz_tab)
        biz_text = QTextEdit()
        biz_text.setReadOnly(True)
        help_lines = []
        help_lines.append("【重要使用说明】")
        help_lines.append("")
        help_lines.append("1. 脱敏输出文档仅作为内部评审草稿，禁止直接对外交付。")
        help_lines.append("2. 对外正式版本需要人工把占位符【■■■】回填真实业务信息。")
        help_lines.append("3. 支持 docx/pdf/xlsx/txt 格式，支持拖拽添加文件。")
        help_lines.append("4. 处理对象为业务合同、采购文件。")
        help_lines.append("5. 处理完成后自动运行可读性复检。")
        help_lines.append("6. 可在脱敏类型面板中按类别控制脱敏/保留。")
        help_lines.append("7. 新增[网址]类别，支持 http/https/www 链接脱敏。")
        help_lines.append("8. 自定义词条面板支持输入框直接添加，回车或点+号添加。")
        help_lines.append("9. 快捷键：Ctrl+O 添加文件  Ctrl+P 开始处理  Ctrl+E 打开输出目录")
        help_lines.append("10. 敏感度说明：")
        help_lines.append("    保守：仅脱敏人名、身份证、电话、银行账号、网址")
        help_lines.append("    标准：脱敏所有黑名单类别（默认）")
        help_lines.append("    激进：脱敏所有黑名单类别，包括地址、邮编、合同编号等")
        help_lines.append("11. 词条前缀说明：")
        help_lines.append("    [W] 白名单（绝对不脱敏）")
        help_lines.append("    [B] 黑名单（强制脱敏）")
        help_lines.append("    无前缀：自动识别")
        biz_text.setPlainText(NL.join(help_lines))
        biz_layout.addWidget(biz_text)
        right_tabs.addTab(biz_tab, "使用说明")

        right_layout.addWidget(right_tabs)
        main_splitter.addWidget(left_panel)
        main_splitter.addWidget(right_panel)
        main_splitter.setStretchFactor(0, 3)
        main_splitter.setStretchFactor(1, 1)
        root.addWidget(main_splitter)

    # ===== 补丁 #4：关于 / 检查更新 / 辅助方法 =====

    def open_about(self):
        """显示关于对话框"""
        msg = (
            "<h2>{name} {ver}</h2>"
            "<p><b>作者：</b>{author}</p>"
            "<p>一款专业的文档智能脱敏工具，支持 Word/Excel/PDF/TXT 格式。</p>"
            "<p><b>功能特性：</b></p>"
            "<ul>"
            "<li>多格式支持：docx / xlsx / pdf / txt</li>"
            "<li>13 种脱敏类型可控</li>"
            "<li>白名单 / 黑名单 / 自定义词条</li>"
            "<li>三档敏感度：保守 / 标准 / 激进</li>"
            "<li>拖拽添加文件，批量处理</li>"
            "<li>可读性自动复检与报告生成</li>"
            "</ul>"
            "<hr>"
            "<p style='color:#666;font-size:8pt;'>"
            "⚠️ 脱敏输出文档仅作为内部评审草稿，禁止直接对外交付。<br/>"
            "对外正式版本需要人工把占位符【■■■】回填真实业务信息。</p>"
        ).format(name=APP_NAME, ver=APP_VERSION, author=APP_AUTHOR)
        QMessageBox.about(self, "关于 " + APP_NAME, msg)

    def check_update(self):
        """检查版本更新（补丁 #4）"""
        import urllib.request
        try:
            req = urllib.request.Request(UPDATE_URL, headers={"User-Agent": APP_NAME + "/" + APP_VERSION})
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            latest_ver = data.get("version", "")
            download_url = data.get("url", "")
            changelog = data.get("changelog", "")
            if latest_ver and self._is_newer(latest_ver, APP_VERSION):
                reply = QMessageBox.question(
                    self, "发现新版本",
                    "当前版本：<b>%s</b><br/>最新版本：<b>%s</b><br/><br/>%s<br/><br/>是否前往下载？"
                    % (APP_VERSION, latest_ver, changelog),
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes,
                )
                if reply == QMessageBox.Yes:
                    import webbrowser
                    if download_url:
                        webbrowser.open(download_url)
                    else:
                        webbrowser.open(UPDATE_URL.rsplit("/", 1)[0])
            else:
                QMessageBox.information(self, "检查更新", "当前已是最新版本（%s）。" % APP_VERSION)
        except Exception as e:
            QMessageBox.warning(self, "检查更新", "无法连接更新服务器：%s" % str(e))

    @staticmethod
    def _is_newer(latest, current):
        """比较两个版本号，判断 latest 是否比 current 新（补丁 #4）"""
        def _parse(v):
            parts = re.sub(r"[^0-9.]", "", v).split(".")
            return [int(p) for p in parts if p.isdigit()]
        la, cur = _parse(latest), _parse(current)
        return la > cur

    def _open_config_dir(self):
        """打开配置文件所在目录"""
        cfg_path = _get_config_dir()
        if os.path.exists(cfg_path):
            os.startfile(cfg_path)
        else:
            QMessageBox.information(self, "提示", "配置目录尚未创建，首次保存设置后将自动生成。")

    # ===== 原有方法（保持不变）=====

    def _add_wl_from_input(self):
        text = self.wl_input.text().strip()
        if text:
            self.wl_list.addItem(QListWidgetItem(text))
            self.config.custom_whitelist.append(text)
            self.config.save()
            self.wl_input.clear()
            self.log("已添加白名单词条：" + text)

    def _del_wl_selected(self):
        for item in self.wl_list.selectedItems():
            row = self.wl_list.row(item)
            self.wl_list.takeItem(row)
            if item.text() in self.config.custom_whitelist:
                self.config.custom_whitelist.remove(item.text())
        self.config.save()

    def _add_bl_from_input(self):
        text = self.bl_input.text().strip()
        if text:
            self.bl_list.addItem(QListWidgetItem(text))
            self.config.custom_blacklist.append(text)
            self.config.save()
            self.bl_input.clear()
            self.log("已添加黑名单词条：" + text)

    def _del_bl_selected(self):
        for item in self.bl_list.selectedItems():
            row = self.bl_list.row(item)
            self.bl_list.takeItem(row)
            if item.text() in self.config.custom_blacklist:
                self.config.custom_blacklist.remove(item.text())
        self.config.save()

    def _on_sens_change(self, text):
        self.config.sensitivity = text
        self.config.save()

    def _on_type_check_changed(self):
        for name, cb in self.type_checks.items():
            self.config.enabled_types[name] = cb.isChecked()
        self.config.save()

    def _show_file_context_menu(self, pos):
        menu = QMenu()
        ra = QAction("移除文件", self)
        ra.triggered.connect(self._remove_selected_file)
        menu.addAction(ra)
        of = QAction("打开文件所在位置", self)
        of.triggered.connect(self._open_file_location)
        menu.addAction(of)
        menu.exec_(self.table.mapToGlobal(pos))

    def _remove_selected_file(self):
        rows = set()
        for item in self.table.selectedItems():
            rows.add(item.row())
        for row in sorted(rows, reverse=True):
            self.table.removeRow(row)

    def _open_file_location(self):
        for item in self.table.selectedItems():
            path = self.table.item(item.row(), 4).text()
            if os.path.exists(path):
                os.startfile(os.path.dirname(path))

    def _open_output_dir(self):
        d = self.output_dir.text().strip()
        if d and os.path.exists(d):
            os.startfile(d)
            return
        for item in self.table.selectedItems():
            path = self.table.item(item.row(), 4).text()
            if os.path.exists(path):
                src = os.path.dirname(path)
                candidates = [f for f in os.listdir(src) if f.startswith("tuomin_task_") and os.path.isdir(os.path.join(src, f))]
                if candidates:
                    latest = max(candidates, key=lambda x: os.path.getmtime(os.path.join(src, x)))
                    out = os.path.join(src, latest)
                    self.output_dir.setText(out)
                    self.config.last_output_dir = out
                    self.config.save()
                    os.startfile(out)
                    return
                os.startfile(src)
                return
        self.log("输出目录不存在或未设置，请先选择输出目录或添加文件")

    def _update_output_dir(self, path):
        self.output_dir.setText(path)
        self.config.last_output_dir = path
        self.config.save()

    def _try_update_output_dir(self, path):
        """尝试从源文件路径推断并更新输出目录"""
        out_dir = self._get_out_dir(path)
        self._update_output_dir(out_dir)

    def _export_custom_dict(self):
        path, _ = QFileDialog.getSaveFileName(self, "导出词条", "custom_dict.txt", "文本文件 (*.txt)")
        if path:
            lines = []
            for i in range(self.wl_list.count()):
                lines.append("[W]" + self.wl_list.item(i).text())
            for i in range(self.bl_list.count()):
                lines.append("[B]" + self.bl_list.item(i).text())
            with open(path, "w", encoding="utf-8") as f:
                f.write(NL.join(lines))
            self.log("词条已导出：" + path)

    def on_add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择文档", "",
            "支持格式 (*.docx *.doc *.pdf *.xlsx *.txt);;Word (*.docx *.doc);;Excel (*.xlsx);;PDF (*.pdf);;文本 (*.txt)"
        )
        for f in files:
            self._add_single_file(f)

    def _add_single_file(self, path):
        if not os.path.exists(path):
            return
        ext = os.path.splitext(path)[1].lower()
        ftype = "Word" if ext in (".docx", ".doc") else ("Excel" if ext == ".xlsx" else ("PDF" if ext == ".pdf" else "文本"))
        size = os.path.getsize(path)
        size_str = self._format_size(size)
        row = self.table.rowCount()
        self.table.insertRow(row)
        self.table.setItem(row, 0, QTableWidgetItem(os.path.basename(path)))
        self.table.setItem(row, 1, QTableWidgetItem(ftype))
        self.table.setItem(row, 2, QTableWidgetItem(size_str))
        self.table.setItem(row, 3, QTableWidgetItem("待处理"))
        self.table.setItem(row, 4, QTableWidgetItem(path))

    def on_clear_list(self):
        self.table.setRowCount(0)
        self.log("已清空文件列表")

    def on_choose_dir(self):
        d = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if d:
            self.output_dir.setText(d)
            self.config.last_output_dir = d
            self.config.save()
            self.log("输出目录：" + d)

    def on_load_dict_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择词典文件", "", "文本文件 (*.txt)")
        if file_path:
            with open(file_path, "r", encoding="utf-8") as f:
                self.wl_list.clear()
                self.bl_list.clear()
                self.config.custom_whitelist = []
                self.config.custom_blacklist = []
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    if line.startswith("[W]"):
                        w = line[3:].strip()
                        if w:
                            self.config.custom_whitelist.append(w)
                            self.wl_list.addItem(QListWidgetItem(w))
                    elif line.startswith("[B]"):
                        b = line[3:].strip()
                        if b:
                            self.config.custom_blacklist.append(b)
                            self.bl_list.addItem(QListWidgetItem(b))
                self.config.save()
            self.log("已从文件加载词典：" + file_path)

    def on_apply_custom_dict(self):
        self.config.custom_whitelist = []
        self.config.custom_blacklist = []
        for i in range(self.wl_list.count()):
            self.config.custom_whitelist.append(self.wl_list.item(i).text())
        for i in range(self.bl_list.count()):
            self.config.custom_blacklist.append(self.bl_list.item(i).text())
        self.custom_dict = list(self.config.custom_blacklist)
        self.config.save()
        self.log("已应用自定义词典：白名单 %d 条，黑名单 %d 条" % (len(self.config.custom_whitelist), len(self.custom_dict)))

    def on_start(self):
        if self._processing:
            self.log("正在处理中，请等待……")
            return
        mode = "脱敏" if self.radio_de.isChecked() else "还原"
        rows = self.table.rowCount()
        if rows == 0:
            QMessageBox.information(self, "提示", "请先添加要处理的文档。")
            return
        self.on_apply_custom_dict()
        self._processing = True
        self.btn_process.setEnabled(False)
        self.progress.setValue(0)
        self.log("开始【%s】处理，共 %d 个文件……" % (mode, rows))
        files = [self.table.item(i, 4).text() for i in range(rows)]
        t = threading.Thread(target=self._process_files, args=(files, mode), daemon=True)
        t.start()

    def _process_files(self, files, mode):
        try:
            total = len(files)
            for idx, f in enumerate(files):
                self._log_safe("正在处理：" + os.path.basename(f))
                ok = self._process_desensitize(f) if mode == "脱敏" else self._process_restore(f)
                status = "完成" if ok else "失败"
                self._update_status(f, status)
                QMetaObject.invokeMethod(self.progress, "setValue", Qt.QueuedConnection, Q_ARG(int, int((idx + 1) / total * 100)))
            self._log_safe("全部处理完成！")
        except Exception as e:
            self._log_safe("处理出错：" + str(e))
            import traceback
            self._log_safe(traceback.format_exc())
        finally:
            self._processing = False
            QMetaObject.invokeMethod(self.btn_process, "setEnabled", Qt.QueuedConnection, Q_ARG(bool, True))

    def _process_desensitize(self, path):
        try:
            parser = DocParser()
            parser.load(path)
            paragraphs = parser.get_paragraphs()
            self._log_safe("  已解析文档，共 %d 个文本片段" % len(paragraphs))
            all_custom = list(self.config.custom_blacklist) + list(self.config.custom_whitelist)
            enabled_types = {name: cb.isChecked() for name, cb in self.type_checks.items()}
            de = Desensitizer(custom_dict=all_custom, enabled_types=enabled_types, sensitivity=self.sens_combo.currentText())
            self._log_safe("  正在分析全文，建立敏感信息映射表……")
            full_text = NL.join(p["text"] for p in paragraphs if p["text"].strip())
            de.build_mapping(full_text)
            replace_items = de.get_replace_items()
            self._log_safe("  " + de.get_mapping_summary())
            if not replace_items:
                self._log_safe("  未识别到敏感信息，将生成原样副本")
                texts = [p["text"] for p in paragraphs]
                self._save_desensitized(path, parser, texts, de, full_text)
                return True
            self._log_safe("  共识别 %d 个敏感实体" % len(replace_items))
            texts = [de.desensitize(p["text"]) for p in paragraphs]
            self._save_desensitized(path, parser, texts, de, full_text)
            return True
        except Exception as e:
            self._log_safe("  脱敏失败：" + str(e))
            self._try_update_output_dir(path)
            import traceback
            self._log_safe("  " + traceback.format_exc())
            return False

    def _save_desensitized(self, path, parser, texts, de, full_text, updated_items=None):
        issues = de.get_readability_issues()
        if issues:
            for issue in issues:
                self._log_safe("  " + issue)
        out_text = NL.join(texts)
        alerts = de.run_validation(full_text, out_text)
        if alerts:
            for alert in alerts:
                self._log_safe("  " + alert)
        else:
            self._log_safe("  校验通过，未发现异常")
        out_dir = self._get_out_dir(path)
        base = os.path.splitext(os.path.basename(path))[0]
        mapping_path = os.path.join(out_dir, "mapping_%s.json" % base)
        de.mapping.save_to_file(mapping_path)
        self._log_safe("  映射表已保存：" + mapping_path)
        ext = os.path.splitext(path)[1].lower()
        if ext == ".xlsx":
            out_path = os.path.join(out_dir, "脱敏_" + base + ".xlsx")
            try:
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                for text in texts:
                    ws.append([text])
                wb.save(out_path)
            except ImportError:
                out_path = os.path.join(out_dir, "脱敏_" + base + ".txt")
                parser.save_txt(out_path, texts)
        elif ext == ".pdf":
            out_path = os.path.join(out_dir, "脱敏_" + base + ".docx")
            parser.save_docx(out_path, texts)
        else:
            out_path = os.path.join(out_dir, "脱敏_" + base + ext)
            if ext == ".txt":
                parser.save_txt(out_path, texts)
            else:
                parser.save_docx(out_path, texts)
        self._log_safe("  脱敏文档已保存：" + out_path)
        self._update_output_dir(out_dir)
        self._generate_report(out_dir, base, de, updated_items)
        self._log_safe("  【重要】脱敏输出文档仅作为内部评审草稿，禁止直接对外交付")
        self._log_safe("  【重要】对外正式版本需要人工把占位符【■■■】回填真实业务信息")

    def _generate_report(self, out_dir, base, de, updated_items=None):
        try:
            lines = []
            lines.append("=" * 50)
            lines.append("文档脱敏处理报告")
            lines.append("=" * 50)
            lines.append("处理时间：%s" % datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            lines.append("敏感度模式：%s" % self.sens_combo.currentText())
            lines.append("")
            lines.append("-- 替换统计 --")
            items = de.get_replace_items()
            if updated_items:
                items = updated_items
            stats = {}
            for item in items:
                if item.enabled:
                    stats[item.category] = stats.get(item.category, 0) + 1
            for cat, cnt in sorted(stats.items()):
                lines.append("  %s: %d处" % (cat, cnt))
            lines.append("")
            lines.append("-- 完整替换映射表 --")
            for orig, ph in sorted(de.mapping.get_all_forward().items(), key=lambda x: -len(x[0])):
                lines.append("  %s -> %s" % (orig, ph))
            lines.append("")
            lines.append("-- 可读性复检 --")
            issues = de.get_readability_issues()
            if issues:
                for issue in issues:
                    lines.append("  " + issue)
            else:
                lines.append("  无异常")
            lines.append("")
            lines.append("-- 注意事项 --")
            lines.append("1. 脱敏输出文档仅作为内部评审草稿，禁止直接对外交付")
            lines.append("2. 对外正式版本需要人工把占位符【■■■】回填真实业务信息")
            lines.append("=" * 50)
            report_path = os.path.join(out_dir, "报告_%s.txt" % base)
            with open(report_path, "w", encoding="utf-8") as f:
                f.write(NL.join(lines))
            self._log_safe("  脱敏报告已保存：" + report_path)
        except Exception as e:
            self._log_safe("  生成报告失败：" + str(e))

    def _process_restore(self, path):
        try:
            parser = DocParser()
            parser.load(path)
            paragraphs = parser.get_paragraphs()
            self._log_safe("  已解析文档，共 %d 个文本片段" % len(paragraphs))
            restorer = Restorer()
            mapping_file = restorer.find_mapping_file(path)
            if not mapping_file:
                self._log_safe("  未找到对应的映射表文件")
                return False
            restorer.load_mapping(mapping_file)
            self._log_safe("  已加载映射表，共 %d 条记录" % restorer.mapping.count)
            texts = [restorer.restore(p["text"]) for p in paragraphs]
            out_dir = self._get_out_dir(path)
            base = os.path.splitext(os.path.basename(path))[0]
            if base.startswith("脱敏_"):
                base = base[3:]
            ext = os.path.splitext(path)[1].lower()
            out_path = os.path.join(out_dir, "还原_" + base + ext)
            if ext == ".txt":
                parser.save_txt(out_path, texts)
            else:
                parser.save_docx(out_path, texts)
            self._log_safe("  还原文档已保存：" + out_path)
            return True
        except Exception as e:
            self._log_safe("  还原失败：" + str(e))
            return False

    def _get_out_dir(self, path):
        custom = self.output_dir.text().strip()
        if custom:
            return custom
        src = os.path.dirname(path)
        folder = "tuomin_task_" + datetime.now().strftime("%Y%m%d_%H%M%S")
        out = os.path.join(src, folder)
        os.makedirs(out, exist_ok=True)
        return out

    def _update_status(self, file_path, status):
        for i in range(self.table.rowCount()):
            if self.table.item(i, 4).text() == file_path:
                self.table.item(i, 3).setText(status)
                break

    def _log_safe(self, msg):
        self.log_signal.append.emit(msg)

    def _append_log(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_panel.appendPlainText("[%s] %s" % (ts, msg))

    def log(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_panel.appendPlainText("[%s] %s" % (ts, msg))

    @staticmethod
    def _format_size(s):
        if s < 1024:
            return "%d B" % s
        elif s < 1024 ** 2:
            return "%.1f KB" % (s / 1024)
        else:
            return "%.1f MB" % (s / 1024 / 1024)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle(QStyleFactory.create("Fusion"))
    app.setStyleSheet(STYLESHEET)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())
